"""Build concise, public Alpha v1.2 financial DOCX documents."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY, GREEN, GREEN_DARK = "17324D", "087F5B", "056246"
GOLD, INK, MUTED = "B7791F", "17212B", "5E6B78"
LIGHT, PALE_GREEN, PALE_GOLD = "F2F4F7", "E8F5EF", "FFF6DD"
WHITE, TABLE_WIDTH, TABLE_INDENT = "FFFFFF", 9360, 120

MODELS = {
    "fidlot": {
        "version": "v5.9", "title_en": "Feedlot livestock fattening",
        "title_ru": "Откормочный комплекс", "cycles": 7, "duration": 35,
        "investment": "$50,000", "payout": "$82,000", "roi": "64.0%",
        "annualized": "21.9%", "operator_cash": "$96,250",
        "operator_property": "$18,000", "operator_total": "$114,250",
        "capital": [("Initial livestock", "Начальный молодняк", "$20,000"),
                    ("Feedlot facility", "Откормочная база", "$18,000"),
                    ("Working reserve", "Оборотный резерв", "$12,000")],
        "economics": [("Cycles 1–2", "Циклы 1–2", "$50,000", "$30,000", "$9,600", "$15,250"),
                      ("Cycles 3–7", "Циклы 3–7", "$50,000", "$26,500", "$8,480", "$13,150"),
                      ("Completion", "Завершение", "—", "—", "$20,400", "$10,000")],
    },
    "hissar": {
        "version": "VariantB v2.1", "title_en": "Hissar sheep breeding",
        "title_ru": "Разведение гиссарской породы", "cycles": 6, "duration": 36,
        "investment": "$50,000", "payout": "$81,672", "roi": "63.3%",
        "annualized": "21.1%", "operator_cash": "$83,160",
        "operator_property": "$18,000", "operator_total": "$101,160",
        "capital": [("38 breeding ewes", "38 овцематок", "$26,600"),
                    ("Sheep shelter", "Кошара", "$18,000"),
                    ("Operating reserve", "Операционный резерв", "$5,400")],
        "economics": [("Cycles 1–2", "Циклы 1–2", "$30,600", "$30,600", "$9,792", "$15,260"),
                      ("Cycles 3–6", "Циклы 3–6", "$30,600", "$24,600", "$10,372", "$11,660"),
                      ("Completion", "Завершение", "—", "—", "$20,600", "$16,000")],
    },
}

TEXT = {
    "en": {
        "role_investor": "Investor model brief",
        "role_operator": "Uzbekistan Feedlot Operator guide",
        "subtitle": "Demonstration economics · 60/40 model · fiat-only Uzbekistan operations",
        "boundary": "Mandatory boundary: the External Investor contracts only with AgriPartners OÜ in Estonia. Approved crypto, if any, stops in Estonia and is converted through approved infrastructure. The Uzbekistan Feedlot Operator receives and returns fiat only under a separate written Operator Agreement.",
        "disclaimer": "Illustrative Alpha v1.2 model. Projections are not guaranteed returns, are not a public offer, and require legal, tax, banking, accounting, compliance, operational, and partner review before any real-money pilot.",
        "metric": "Metric", "value": "Value", "capital": "Illustrative use of initial capital",
        "item": "Item", "amount": "Amount", "economics": "Condensed cycle economics",
        "phase": "Phase", "revenue": "Revenue", "net": "Net pool",
        "investor_cash": "Investor cash", "operator_cash": "Operator cash",
        "roles": "Roles and authoritative evidence", "role": "Role", "responsibility": "Responsibility and evidence",
        "role_rows": [
            ("External Investor", "Contracts with and funds AgriPartners OÜ; no direct transfer to Uzbekistan."),
            ("AgriPartners OÜ", "Legal counterparty; conversion, banking, accounting, reconciliation and governance."),
            ("Uzbekistan Feedlot Operator", "Separate agreement; fiat receipts and returns; operational delivery and evidence."),
            ("Farmer product role", "Operational reports and confirmations only; no wallet, crypto receipt or on-chain payment duty."),
            ("NEAR Testnet", "Demo audit and automation on the Estonia/investor side; never replaces authoritative records."),
        ],
        "controls": "Assumptions, controls and risks", "area": "Area", "position": "Position",
        "control_rows": [
            ("Model", "Illustrative prices, quantities, costs, timing, 60/40 split and 20% performance fee."),
            ("Reserve", "Exploratory protection reserve; not insurance or a guarantee; release may stop after loss, default, missing reports or dispute."),
            ("Evidence", "Signed agreements, bank/payment records, invoices, accounting and reconciliation remain authoritative."),
            ("Risks", "Animal health, mortality, feed and market prices, FX, execution, counterparties, regulation and liquidity."),
        ],
        "before": "Before a real pilot",
        "before_rows": ["Approve company, banking/payment, accounting, compliance and crypto-to-fiat providers.",
                        "Execute separate Investor and Operator agreements after qualified legal review.",
                        "Validate assumptions with a named operator, suppliers, veterinary controls and local evidence.",
                        "Use Alpha v1.2 only as a Testnet workflow demonstration until formal pilot approval."],
        "source": "Source: AgriPartners 60/40 model. Simple annualized ROI equals total ROI divided by model duration; it is not APR or IRR.",
    },
    "ru": {
        "role_investor": "Краткая модель для инвестора",
        "role_operator": "Руководство оператора откормочного комплекса в Узбекистане",
        "subtitle": "Демонстрационная экономика · модель 60/40 · операции в Узбекистане только в фиате",
        "boundary": "Обязательная граница: внешний инвестор заключает договор только с AgriPartners OÜ в Эстонии. Разрешённая криптовалюта, если используется, остаётся в Эстонии и конвертируется через одобренную инфраструктуру. Оператор в Узбекистане получает и возвращает только фиат по отдельному письменному договору.",
        "disclaimer": "Иллюстративная модель Alpha v1.2. Прогнозы не гарантируют доходность, не являются публичным предложением и требуют юридической, налоговой, банковской, бухгалтерской, комплаенс-, операционной и партнёрской проверки до пилота с реальными средствами.",
        "metric": "Показатель", "value": "Значение", "capital": "Иллюстративное использование стартового капитала",
        "item": "Статья", "amount": "Сумма", "economics": "Сокращённая экономика циклов",
        "phase": "Этап", "revenue": "Выручка", "net": "Чистый пул",
        "investor_cash": "Инвестору", "operator_cash": "Оператору",
        "roles": "Роли и официальные подтверждения", "role": "Роль", "responsibility": "Ответственность и подтверждения",
        "role_rows": [
            ("Внешний инвестор", "Договор и финансирование только с AgriPartners OÜ; без прямых переводов в Узбекистан."),
            ("AgriPartners OÜ", "Юридический контрагент; конвертация, банк, учёт, сверка и управление."),
            ("Оператор в Узбекистане", "Отдельный договор; фиатные поступления и возвраты; операции и подтверждающие документы."),
            ("Роль Farmer в продукте", "Только отчётность и подтверждения; без кошелька, получения криптовалюты и ончейн-платежей."),
            ("NEAR Testnet", "Демо-аудит и автоматизация на стороне Эстонии/инвестора; не заменяет официальные документы."),
        ],
        "controls": "Предположения, контроль и риски", "area": "Область", "position": "Положение",
        "control_rows": [
            ("Модель", "Иллюстративные цены, объёмы, затраты, сроки, сплит 60/40 и performance fee 20%."),
            ("Резерв", "Исследовательский механизм защиты; не страховка и не гарантия; выдача может остановиться при убытке, дефолте, отсутствии отчёта или споре."),
            ("Документы", "Подписанные договоры, банковские документы, счета, бухгалтерский учёт и сверка являются официальными источниками."),
            ("Риски", "Здоровье и падёж скота, корма и цены, валютный риск, исполнение, контрагенты, регулирование и ликвидность."),
        ],
        "before": "До пилота с реальными средствами",
        "before_rows": ["Утвердить компанию и банковскую, платёжную, бухгалтерскую, комплаенс- и crypto-to-fiat инфраструктуру.",
                        "После квалифицированной юридической проверки подписать отдельные договоры с инвестором и оператором.",
                        "Проверить допущения с выбранным оператором, поставщиками, ветеринарным контролем и локальными документами.",
                        "Использовать Alpha v1.2 только как Testnet-демонстрацию до формального утверждения пилота."],
        "source": "Источник: модель AgriPartners 60/40. Простой годовой ROI — общий ROI, делённый на срок модели; это не APR и не IRR.",
    },
}


def font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rpr.rFonts.set(qn(key), "Arial")
    run.font.size, run.bold, run.italic = Pt(size), bold, italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
    if node.getparent() is None:
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def geometry(table, widths):
    table.autofit, table.alignment = False, WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr
    for tag, value in (("w:tblW", sum(widths)), ("w:tblInd", TABLE_INDENT)):
        node = pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
        if node.getparent() is None:
            pr.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            pr = cell._tc.get_or_add_tcPr()
            node = pr.find(qn("w:tcW"))
            if node is None:
                node = OxmlElement("w:tcW")
            if node.getparent() is None: pr.append(node)
            node.set(qn("w:w"), str(widths[index])); node.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def configure(doc, label):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.18
    for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7), ("Heading 3", 12, 10, 5)):
        style = doc.styles[name]; style.font.name = "Arial"; style.font.size = Pt(size)
        style.font.bold = True; style.font.color.rgb = RGBColor.from_string(GREEN_DARK)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    font(sec.header.paragraphs[0].add_run(label), 8, MUTED, True)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("AgriPartners · Alpha v1.2"), 8, MUTED)


def para(doc, text, size=10.5, color=INK, bold=False, italic=False, after=6, align=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.18
    if align is not None: p.alignment = align
    font(p.add_run(text), size, color, bold, italic); return p


def heading(doc, text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}"); font(p.add_run(text), 13 if level == 2 else 16, GREEN_DARK, True)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent = Inches(.375)
    p.paragraph_format.first_line_indent = Inches(-.188); p.paragraph_format.space_after = Pt(4)
    font(p.add_run(text), 10.2)


def title(doc, kicker, name, subtitle):
    para(doc, kicker.upper(), 9, GREEN, bold=True, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, name, 25, NAVY, bold=True, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, subtitle, 12.5, MUTED, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)


def callout(doc, text, fill=PALE_GREEN, color=GREEN_DARK):
    table = doc.add_table(rows=1, cols=1); geometry(table, [TABLE_WIDTH]); shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]; p.paragraph_format.space_after = Pt(0); font(p.add_run(text), 9.7, color, True)
    para(doc, "", 1, after=2)


def table(doc, headers, rows, widths, size=8.7, fill=NAVY):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = "Table Grid"; geometry(tbl, widths)
    for cell, value in zip(tbl.rows[0].cells, headers):
        shade(cell, fill); p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0); font(p.add_run(str(value)), size, WHITE, True)
    for row_index, values in enumerate(rows):
        cells = tbl.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            if row_index % 2: shade(cell, LIGHT)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0); font(p.add_run(str(value)), size)


def build(model_key, audience, lang, output):
    m, t = MODELS[model_key], TEXT[lang]
    role = t["role_investor"] if audience == "investor" else t["role_operator"]
    name = m["title_en"] if lang == "en" else m["title_ru"]
    doc = Document(); configure(doc, f"AgriPartners | {role} | {m['version']}")
    title(doc, role, name, f"{m['version']} · {t['subtitle']}")
    callout(doc, t["boundary"]); callout(doc, t["disclaimer"], PALE_GOLD, GOLD)
    metrics = ([
        ("Initial investment" if lang == "en" else "Стартовое финансирование", m["investment"]),
        ("Projected investor payout" if lang == "en" else "Прогноз выплаты инвестору", m["payout"]),
        ("Net ROI" if lang == "en" else "Чистый ROI", m["roi"]),
        ("Simple annualized ROI" if lang == "en" else "Простой годовой ROI", m["annualized"]),
    ] if audience == "investor" else [
        ("Projected operator cash" if lang == "en" else "Прогноз выплат оператору", m["operator_cash"]),
        ("Projected property transfer" if lang == "en" else "Прогноз передачи имущества", m["operator_property"]),
        ("Projected total benefit" if lang == "en" else "Прогноз общей выгоды", m["operator_total"]),
        ("Cycles / duration" if lang == "en" else "Циклы / срок", f"{m['cycles']} / {m['duration']} mo."),
    ])
    table(doc, [t["metric"], t["value"]], metrics, [5900, 3460], 9.2, GREEN_DARK)
    heading(doc, t["capital"])
    capital = [(r[0] if lang == "en" else r[1], r[2]) for r in m["capital"]] + [("Total" if lang == "en" else "Итого", m["investment"])]
    table(doc, [t["item"], t["amount"]], capital, [6600, 2760])
    heading(doc, t["economics"])
    econ = [(r[0] if lang == "en" else r[1], *r[2:]) for r in m["economics"]]
    table(doc, [t["phase"], t["revenue"], t["net"], t["investor_cash"], t["operator_cash"]], econ, [2200, 1700, 1750, 1800, 1910], 8.0)

    doc.add_page_break(); title(doc, "AgriPartners", t["roles"], f"{name} · {m['version']}")
    table(doc, [t["role"], t["responsibility"]], t["role_rows"], [2500, 6860], 8.7, GREEN_DARK)
    heading(doc, t["controls"]); table(doc, [t["area"], t["position"]], t["control_rows"], [1800, 7560], 8.6)
    heading(doc, t["before"])
    for item in t["before_rows"]: bullet(doc, item)
    para(doc, t["source"], 8.5, MUTED, False, True, 4); callout(doc, t["disclaimer"], PALE_GOLD, GOLD)
    doc.core_properties.title = f"AgriPartners {role} - {name} - {lang.upper()}"
    doc.core_properties.subject = "Concise Alpha v1.2 demonstration financial model"
    doc.core_properties.author = "AgriPartners"
    output.parent.mkdir(parents=True, exist_ok=True); doc.save(output)


def main():
    parser = argparse.ArgumentParser(); parser.add_argument("--root", default=str(Path(__file__).resolve().parents[1]))
    root = Path(parser.parse_args().root).resolve(); base = root / "docs" / "60-40" / "source"
    for lang in ("en", "ru"):
        for audience in ("investor", "operator"):
            for key in ("fidlot", "hissar"):
                label = "Fidlot-v5.9" if key == "fidlot" else "VariantB-v2.1"
                role = "Investor" if audience == "investor" else "Operator"
                output = base / lang / f"Agri-{role}-{label}-6040-{lang.upper()}.docx"
                build(key, audience, lang, output); print(output)


if __name__ == "__main__": main()
